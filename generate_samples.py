from fpdf import FPDF
from docx import Document
import os

data_dir = "data"
if not os.path.exists(data_dir):
    os.makedirs(data_dir)

# 1. Create a sample PDF file
def create_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    content = """Service Level Agreement (SLA): Enterprise Fiber
    
1. Uptime Guarantee
We guarantee a 99.99% uptime for all Enterprise Fiber connections.

2. Mean Time To Repair (MTTR)
In the event of a total service disruption, the target MTTR is 4 hours from the time the ticket is acknowledged.

3. SLA Credits
If the uptime falls below 99.99% in a given month, the customer is eligible for a 5% credit on their monthly bill.
"""
    for line in content.split("\n"):
        pdf.cell(200, 10, txt=line, ln=True, align="L")
        
    pdf_path = os.path.join(data_dir, "sla_document.pdf")
    pdf.output(pdf_path)
    print(f"Created {pdf_path}")

# 2. Create a sample DOCX file
def create_docx():
    doc = Document()
    doc.add_heading("Hardware Specifications: 5G Base Station", 0)
    
    doc.add_paragraph(
        "This document details the hardware specifications for the new 5G base station deployments."
    )
    
    doc.add_heading("Power Requirements", level=1)
    doc.add_paragraph("Input voltage: -48V DC nominal (-36V to -60V operational range).")
    doc.add_paragraph("Maximum power consumption under full load is 1500W.")
    
    doc.add_heading("Cooling System", level=1)
    doc.add_paragraph("The base station uses active fan cooling. Fans must be replaced every 3 years to maintain optimal performance.")
    
    docx_path = os.path.join(data_dir, "hardware_specs.docx")
    doc.save(docx_path)
    print(f"Created {docx_path}")

if __name__ == "__main__":
    create_pdf()
    create_docx()
